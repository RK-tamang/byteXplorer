from docx import Document
from docx.shared import Inches, Pt
from pathlib import Path
from models.analysis import AnalysisResult
from .base_reporter import BaseReporter

class DOCXReporter(BaseReporter):
    def generate(self, output_path: Path):
        doc = Document()
        
        # Title
        title = doc.add_heading("Static Malware Analysis Report", 0)
        title.runs[0].font.size = Pt(16)
        
        # Executive Summary
        doc.add_heading("Executive Summary", level=1)
        p = doc.add_paragraph(self.result.executive_summary)
        
        # File Info
        doc.add_heading("File Information", level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Attribute"
        hdr_cells[1].text = "Value"
        
        file_info_rows = [
            ("Filename", str(self.result.file_info.filename)),
            ("Full Path", self.result.file_info.full_path),
            ("Size", f"{self.result.file_info.size_bytes:,} bytes"),
            ("File Type", self.result.file_info.file_type),
            ("Is PE", str(self.result.file_info.is_pe)),
        ]
        if self.result.pe_header:
            file_info_rows += [
                ("Machine Type", self.result.pe_header.machine_type),
                ("Image Base", f"0x{self.result.pe_header.image_base:08X}"),
                ("Entry Point", f"0x{self.result.pe_header.entry_point:08X}"),
            ]
        for attr, val in file_info_rows:
            row_cells = table.add_row().cells
            row_cells[0].text = attr
            row_cells[1].text = val
        
        # Hashes
        doc.add_heading("Hashes", level=1)
        doc.add_paragraph(f"MD5: {self.result.hashes.md5}")
        doc.add_paragraph(f"SHA1: {self.result.hashes.sha1}")
        doc.add_paragraph(f"SHA256: {self.result.hashes.sha256}")
        if self.result.hashes.imphash:
            doc.add_paragraph(f"imphash: {self.result.hashes.imphash}")
        
        # Assessment
        doc.add_heading("Risk Assessment", level=1)
        doc.add_paragraph(f"Severity: {self.result.assessment.severity.value}")
        doc.add_paragraph(f"Score: {self.result.assessment.score}")
        doc.add_paragraph(f"Packer Status: {self.result.assessment.packer_status.value}")
        doc.add_paragraph(f"Reasons: {', '.join(self.result.assessment.reasons) or 'None'}")
        
        doc.save(str(output_path) + ".docx")

